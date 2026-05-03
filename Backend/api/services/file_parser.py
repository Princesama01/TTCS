import io
from typing import Tuple

from docx import Document
from pypdf import PdfReader


def parse_docx(file_content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_pdf(file_content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_content))
        pages = []
        for page in reader.pages:
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(text)
        content = "\n\n".join(pages).strip()
        if not content:
            raise ValueError("PDF has no extractable text. It may be scanned/image-based.")
        return content
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_file(file_content: bytes, filename: str) -> Tuple[str, str]:
    filename_lower = filename.lower()
    if filename_lower.endswith(".docx"):
        return parse_docx(file_content), "docx"
    if filename_lower.endswith(".pdf"):
        return parse_pdf(file_content), "pdf"
    raise ValueError(f"Unsupported file type: {filename}")
